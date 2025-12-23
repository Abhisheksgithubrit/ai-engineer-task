import streamlit as st
import os
import json
import tempfile
import pandas as pd
from docx import Document
from docx.shared import RGBColor   # ✅ FIX ADDED
import re
import base64
from io import BytesIO

# ---------------- PAGE CONFIG ----------------
st.set_page_config(
    page_title="ADGM-Compliant Corporate Agent",
    page_icon="📝",
    layout="wide"
)

# ---------------- DOCUMENT REQUIREMENTS ----------------
DOCUMENT_REQUIREMENTS = {
    "Company Formation": [
        "Articles of Association (AoA)",
        "Memorandum of Association (MoA/MoU)",
        "Board Resolution Templates",
        "Shareholder Resolution Templates",
        "Incorporation Application Form",
        "UBO Declaration Form",
        "Register of Members and Directors",
        "Change of Registered Address Notice"
    ]
}

# ---------------- ADGM REGULATIONS ----------------
ADGM_REGULATIONS = {
    "Articles of Association": [
        {
            "pattern": r"UAE Federal",
            "issue": "Incorrect legal framework (ADGM has independent laws)",
            "severity": "High",
            "regulation": "ADGM Companies Regulations 2020"
        },
        {
            "pattern": r"AED|Dirham",
            "issue": "Share capital must be denominated in USD",
            "severity": "Medium",
            "regulation": "ADGM Companies Regulations 2020"
        },
        {
            "pattern": r"Dubai Courts",
            "issue": "Incorrect jurisdiction (should be ADGM Courts)",
            "severity": "High",
            "regulation": "ADGM Courts Regulations 2015"
        }
    ],
    "General": [
        {
            "pattern": r"Dubai Courts",
            "issue": "Incorrect jurisdiction (should be ADGM Courts)",
            "severity": "High",
            "regulation": "ADGM Courts Regulations 2015"
        }
    ]
}

# ---------------- CLASSIFIER ----------------
def create_document_type_classifier():
    return {
        "Articles of Association (AoA)": {
            "keywords": ["articles", "association"],
            "category": "Company Formation"
        }
    }

def classify_document(text, classifier):
    text = text.lower()
    for doc_type, info in classifier.items():
        if any(k in text for k in info["keywords"]):
            return doc_type
    return "Unknown Document"

# ---------------- RED FLAG DETECTION ----------------
def detect_red_flags(text, doc_type):
    issues = []
    rules = ADGM_REGULATIONS.get(doc_type, []) + ADGM_REGULATIONS.get("General", [])

    for rule in rules:
        for match in re.finditer(rule["pattern"], text, re.IGNORECASE):
            issues.append({
                "document": doc_type,
                "issue": rule["issue"],
                "severity": rule["severity"],
                "suggestion": rule["regulation"],
                "matched_text": match.group(),
                "position": match.start()
            })
    return issues

# ---------------- INSERT COMMENTS (FIXED) ----------------
def insert_comments(doc, issues):
    issues = sorted(issues, key=lambda x: x["position"], reverse=True)

    for issue in issues:
        pos = issue["position"]
        count = 0

        for para in doc.paragraphs:
            if count <= pos < count + len(para.text):
                rel = pos - count
                text = para.text
                para.text = ""

                if rel > 0:
                    para.add_run(text[:rel])

                flagged = para.add_run(text[rel:rel + len(issue["matched_text"])])
                flagged.font.color.rgb = RGBColor(255, 0, 0)  # 🔴 RED FIXED

                comment = para.add_run(
                    f" [ISSUE: {issue['issue']} | {issue['suggestion']}]"
                )
                comment.font.color.rgb = RGBColor(0, 0, 255)  # 🔵 BLUE FIXED
                comment.font.italic = True

                if rel + len(issue["matched_text"]) < len(text):
                    para.add_run(text[rel + len(issue["matched_text"]):])
                break

            count += len(para.text) + 1

    return doc

# ---------------- PROCESS DOCUMENT ----------------
def process_uploaded_documents(files):
    classifier = create_document_type_classifier()
    all_docs = []
    all_issues = []

    for f in files:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tmp.write(f.getbuffer())
            path = tmp.name

        doc = Document(path)
        text = "\n".join(p.text for p in doc.paragraphs)

        doc_type = classify_document(text, classifier)
        issues = detect_red_flags(text, doc_type)

        doc = insert_comments(doc, issues)

        output = BytesIO()
        doc.save(output)
        output.seek(0)

        all_docs.append({
            "name": f.name,
            "type": doc_type,
            "file": output,
            "issues": len(issues)
        })

        all_issues.extend(issues)
        os.unlink(path)

    return all_docs, all_issues

# ---------------- STREAMLIT UI ----------------
def main():
    st.title("ADGM-Compliant Corporate Agent")

    uploaded_files = st.file_uploader(
        "Upload ADGM-related documents (.docx)",
        type=["docx"],
        accept_multiple_files=True
    )

    if not uploaded_files:
        st.info("Upload a .docx file to begin.")
        return

    docs, issues = process_uploaded_documents(uploaded_files)

    st.header("Document Review Results")

    col1, col2 = st.columns(2)
    col1.metric("Total Issues", len(issues))
    col2.metric("High Severity", len([i for i in issues if i["severity"] == "High"]))

    st.subheader("Documents Analyzed")
    for d in docs:
        with st.expander(f"{d['name']} ({d['issues']} issues)"):
            st.markdown(
                f'<a download="Reviewed_{d["name"]}" href="data:application/octet-stream;base64,{base64.b64encode(d["file"].getvalue()).decode()}">Download Reviewed Document</a>',
                unsafe_allow_html=True
            )

            if d["issues"] == 0:
                st.success("No issues found.")
            else:
                for i in issues:
                    st.markdown(f"🔴 **{i['severity']}** – {i['issue']}")

if __name__ == "__main__":
    main()
